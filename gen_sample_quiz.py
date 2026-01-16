import docx
from docx.shared import RGBColor

def create_sample():
    doc = docx.Document()
    doc.add_heading('Sample Quiz for Import', 0)
    doc.add_paragraph('Import Instructions: Upload this file to test the import functionality.')

    # Question 1: Using Bold for correct answer
    doc.add_paragraph('1. What is the capital of France?')
    doc.add_paragraph('a. London')
    doc.add_paragraph('b. Berlin')
    
    # Correct answer (Bold)
    p = doc.add_paragraph()
    run = p.add_run('c. Paris')
    run.bold = True
    
    doc.add_paragraph('d. Madrid')
    doc.add_paragraph('')

    # Question 2: Using Red Color for correct answer
    doc.add_paragraph('2. Which web framework is written in Python?')
    doc.add_paragraph('A. Laravel')
    doc.add_paragraph('B. Spring')
    
    # Correct answer (Red)
    p = doc.add_paragraph()
    run = p.add_run('C. Django')
    run.font.color.rgb = RGBColor(255, 0, 0)

    # Question 3: Simple Bold
    doc.add_paragraph('3. 2 + 2 = ?')
    p = doc.add_paragraph()
    p.add_run('a. 4').bold = True
    doc.add_paragraph('b. 5')
    doc.add_paragraph('c. 6')

    doc.save('sample_quiz_for_import.docx')
    print("File 'sample_quiz_for_import.docx' created successfully.")

if __name__ == "__main__":
    create_sample()
